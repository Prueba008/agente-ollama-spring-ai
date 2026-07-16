from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
DIAGRAMS = OUT / "diagramas"
OUT.mkdir(exist_ok=True)
DIAGRAMS.mkdir(exist_ok=True)

BLUE = "1F4E78"
MID_BLUE = "2E75B6"
LIGHT_BLUE = "D9EAF7"
PALE = "F3F6F9"
GRAY = "5B6573"
INK = "17212B"
GREEN = "2E7D32"
RED = "B3261E"
GOLD = "9A6700"


def font(size, bold=False):
    candidates = ["/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                  "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"]
    if bold:
        candidates = ["/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                      "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def box(draw, xy, title, body="", fill="#F3F6F9", outline="#2E75B6"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=4)
    draw.text((x1 + 22, y1 + 16), title, font=font(27, True), fill="#17212B")
    if body:
        lines = body.split("\n")
        y = y1 + 58
        for line in lines:
            draw.text((x1 + 22, y), line, font=font(20), fill="#44515E")
            y += 29


def arrow(draw, start, end, label="", color="#2E75B6"):
    draw.line([start, end], fill=color, width=6)
    x2, y2 = end
    x1, y1 = start
    import math
    angle = math.atan2(y2-y1, x2-x1)
    length = 20
    for delta in (2.55, -2.55):
        p = (x2 + length * math.cos(angle + delta), y2 + length * math.sin(angle + delta))
        draw.line([end, p], fill=color, width=6)
    if label:
        mx, my = (x1+x2)//2, (y1+y2)//2
        label_font = font(17)
        bounds = draw.textbbox((0,0), label, font=label_font)
        w = bounds[2] - bounds[0]
        draw.rounded_rectangle((mx-w/2-12, my-19, mx+w/2+12, my+19), 9, fill="white")
        draw.text((mx-w/2, my-11), label, font=label_font, fill="#44515E")


def canvas(title, subtitle):
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1600, 105), fill="#1F4E78")
    d.text((55, 22), title, font=font(34, True), fill="white")
    d.text((56, 66), subtitle, font=font(19), fill="#D9EAF7")
    return img, d


def diagram_architecture():
    img, d = canvas("Arquitectura de validación semántica", "Separación entre señal léxica, juez semántico y referencia humana")
    box(d, (70, 175, 400, 355), "Respuesta del agente", "afirmación\nfuentes\ngrounded\nmodelConfidence", "#F3F6F9")
    box(d, (70, 535, 400, 715), "Evidencia RAG", "chunks recuperados\nfuentes válidas\ncontexto disponible", "#F3F6F9")
    box(d, (560, 175, 930, 355), "SimpleEvaluator", "solapamiento léxico\nrelevancia\nfaithfulness aproximada", "#FFF4D6", "#9A6700")
    box(d, (560, 535, 930, 715), "SemanticJudge", "NLI o LLM-as-judge\nENTAILED\nCONTRADICTED / NEI", "#D9EAF7")
    box(d, (1110, 175, 1515, 355), "Métricas operativas", "riesgo y controles\nno equivalen a verdad\nni probabilidad calibrada", "#FDE7E5", "#B3261E")
    box(d, (1110, 535, 1515, 715), "Benchmark humano", "matriz de confusión\naccuracy / precision\nrecall / F1", "#E5F4E7", "#2E7D32")
    arrow(d, (400, 265), (560, 265), "texto")
    arrow(d, (400, 625), (560, 625), "evidencia")
    arrow(d, (745, 535), (745, 355), "contraste")
    arrow(d, (930, 265), (1110, 265), "señales")
    arrow(d, (930, 625), (1110, 625), "etiquetas")
    d.text((70, 815), "Regla central: modelConfidence nunca decide por sí sola si una afirmación es verdadera.", font=font(24, True), fill="#B3261E")
    path = DIAGRAMS / "01_arquitectura_validacion_semantica.png"
    img.save(path, quality=95)
    return path


def diagram_sequence():
    img, d = canvas("Secuencia de evaluación", "Ejecución por caso y consolidación contra gold set humano")
    xs = [140, 500, 860, 1220]
    names = ["Benchmark", "SemanticJudge", "Corpus humano", "Resultado"]
    for x, name in zip(xs, names):
        d.rounded_rectangle((x-105, 140, x+105, 205), 14, fill="#D9EAF7", outline="#2E75B6", width=3)
        tw = d.textbbox((0,0), name, font=font(22, True))[2]
        d.text((x-tw/2, 158), name, font=font(22, True), fill="#17212B")
        d.line((x, 205, x, 810), fill="#A9B4BF", width=3)
    events = [
        (260, 140, 860, "1. leer evidencia, claim y etiqueta humana"),
        (350, 140, 500, "2. judge(evidence, claim)"),
        (440, 500, 140, "3. etiqueta predicha"),
        (530, 140, 860, "4. comparar predicha vs. humana"),
        (620, 140, 1220, "5. actualizar matriz de confusión"),
        (710, 1220, 140, "6. accuracy, precision, recall y F1")]
    for y, x1, x2, label in events:
        arrow(d, (x1, y), (x2, y), label, "#2E75B6" if x2 > x1 else "#2E7D32")
    path = DIAGRAMS / "02_secuencia_evaluacion_semantica.png"
    img.save(path, quality=95)
    return path


def diagram_flow():
    img, d = canvas("Flujo de decisión y calibración", "Cómo interpretar una afirmación sin confundir confianza con certeza")
    box(d, (570, 135, 1030, 245), "Afirmación + evidencia", fill="#D9EAF7")
    box(d, (570, 330, 1030, 455), "¿La evidencia implica, contradice\no no alcanza?", fill="#FFF4D6", outline="#9A6700")
    box(d, (80, 585, 430, 735), "ENTAILED", "respaldada semánticamente\nvalidar fuente y alcance", "#E5F4E7", "#2E7D32")
    box(d, (625, 585, 975, 735), "CONTRADICTED", "no presentar como verdad\nriesgo crítico", "#FDE7E5", "#B3261E")
    box(d, (1170, 585, 1520, 735), "NOT ENOUGH INFO", "pedir evidencia\no reconocer ausencia", "#F3F6F9", "#5B6573")
    arrow(d, (800, 245), (800, 330), "evaluar")
    arrow(d, (650, 455), (255, 585), "implica", "#2E7D32")
    arrow(d, (800, 455), (800, 585), "contradice", "#B3261E")
    arrow(d, (950, 455), (1345, 585), "insuficiente", "#5B6573")
    d.rounded_rectangle((330, 800, 1270, 865), 14, fill="#1F4E78")
    msg = "Contrastar con etiquetas humanas; modelConfidence es solo una señal secundaria."
    msg_font = font(20, True)
    tw = d.textbbox((0,0), msg, font=msg_font)[2]
    d.text((800-tw/2, 821), msg, font=msg_font, fill="white")
    path = DIAGRAMS / "03_flujo_decision_semantica.png"
    img.save(path, quality=95)
    return path


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, (h, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]; cell.width = Inches(width); shade(cell, LIGHT_BLUE); margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(h), 10, True, BLUE)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width); margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2: shade(cells[i], PALE)
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), 9.3, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), 10.5)
    return p


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), 11, True, BLUE)
        set_font(p.add_run(text[len(bold_lead):]), 11)
    else:
        set_font(p.add_run(text), 11)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_figure(doc, path, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(6.35))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(8)
    set_font(c.add_run(caption), 9, False, GRAY)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, 9, False, GRAY)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def build_doc(diagrams):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.4); sec.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in (("Heading 1",16,16,8,MID_BLUE),("Heading 2",13,12,6,MID_BLUE),("Heading 3",12,8,4,BLUE)):
        st = styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    for lname in ("List Bullet", "List Number"):
        st=styles[lname]; st.font.name="Calibri"; st.font.size=Pt(10.5); st.paragraph_format.left_indent=Inches(0.5); st.paragraph_format.first_line_indent=Inches(-0.25); st.paragraph_format.space_after=Pt(4)

    header = sec.header.paragraphs[0]
    set_font(header.add_run("MAASTER.IA  |  VALIDACIÓN SEMÁNTICA"), 9, True, BLUE)
    footer = sec.footer.paragraphs[0]; add_page_number(footer)

    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(35); p.paragraph_format.space_after=Pt(5)
    set_font(p.add_run("INFORME TÉCNICO"), 11, True, MID_BLUE)
    p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(8)
    set_font(p.add_run("Validación de verdad semántica\nen agentes RAG con Ollama y Spring AI"), 25, True, BLUE)
    p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(22)
    set_font(p.add_run("Paráfrasis, negaciones, contradicciones complejas y calibración contra etiquetas humanas"), 13, False, GRAY)
    for label, value in (("Proyecto", "agente-ollama-spring-ai"),("Repositorio", "github.com/Prueba008/agente-ollama-spring-ai"),("Versión", "1.0"),("Fecha", "16 de julio de 2026"),("Alcance", "SemanticJudge, benchmark y pruebas unitarias")):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        set_font(p.add_run(label+": "), 10.5, True, BLUE); set_font(p.add_run(value), 10.5)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(22)
    set_font(p.add_run("Conclusión ejecutiva"), 12, True, MID_BLUE)
    body(doc, "La coincidencia textual no demuestra verdad. La validación implementada separa el solapamiento léxico de la relación semántica y obliga a medir el juez contra un corpus humano etiquetado. La confianza informada por el modelo se conserva como señal auxiliar, nunca como veredicto autónomo.")

    doc.add_page_break()
    heading(doc, "1. Propósito y alcance", 1)
    body(doc, "El objetivo es verificar si una respuesta está implicada por la evidencia recuperada, si la contradice o si la información disponible resulta insuficiente. La suite complementa el evaluador determinístico existente y hace explícitos sus límites frente a fenómenos lingüísticos que conservan vocabulario pero cambian significado.")
    bullet(doc, "Distinguir verdad semántica de similitud léxica.")
    bullet(doc, "Detectar paráfrasis válidas aunque no repitan los mismos términos.")
    bullet(doc, "Detectar negaciones explícitas, implícitas y contradicciones multioración.")
    bullet(doc, "Contrastar NLI o LLM-as-judge contra etiquetas humanas versionadas.")
    bullet(doc, "Evitar que modelConfidence sea interpretada como probabilidad de verdad.")
    heading(doc, "2. Principio de validación", 1)
    body(doc, "La evaluación se expresa como una clasificación de inferencia textual natural. Cada caso contiene evidencia, una afirmación y una etiqueta humana de referencia. El juez recibe solamente evidencia y afirmación; no recibe la confianza autoinformada del agente.")
    add_table(doc, ["Etiqueta", "Significado", "Acción esperada"], [
        ("ENTAILED", "La evidencia implica la afirmación.", "Puede considerarse respaldada, sujeto a fuente y alcance."),
        ("CONTRADICTED", "La evidencia niega o resulta incompatible con la afirmación.", "No presentar como verdad; elevar riesgo."),
        ("NOT_ENOUGH_INFORMATION", "La evidencia no permite resolver la afirmación.", "Reconocer ausencia o recuperar más contexto.")], [1.35, 2.55, 2.6])
    add_figure(doc, diagrams[0], "Figura 1. Componentes y límites de la validación semántica.")

    heading(doc, "3. Componentes implementados", 1)
    add_table(doc, ["Componente", "Responsabilidad"], [
        ("SemanticLabel", "Define ENTAILED, CONTRADICTED y NOT_ENOUGH_INFORMATION."),
        ("SemanticJudge", "Contrato funcional para un modelo NLI o un LLM-as-judge."),
        ("HumanLabeledSemanticCase", "Caso versionable con evidencia, claim y etiqueta humana."),
        ("SemanticJudgeBenchmark", "Ejecuta el juez, construye la matriz y calcula métricas."),
        ("SemanticTruthTest", "Prueba fenómenos lingüísticos y el falso positivo léxico."),
        ("SemanticJudgeBenchmarkTest", "Prueba cálculos, errores de entrada e inmutabilidad.")], [2.1, 4.4])
    heading(doc, "4. Secuencia de ejecución", 1)
    body(doc, "Por cada caso humano, el benchmark solicita una predicción al juez, la compara con la etiqueta de referencia y acumula la celda correspondiente de la matriz de confusión. Una vez procesado el corpus completo, expone métricas por etiqueta.")
    add_figure(doc, diagrams[1], "Figura 2. Secuencia de evaluación y consolidación de métricas.")

    doc.add_page_break()
    heading(doc, "5. Casos de verdad semántica", 1)
    add_table(doc, ["Caso", "Evidencia", "Afirmación", "Etiqueta"], [
        ("Paráfrasis", "Outbox guarda el evento en la transacción del negocio.", "Evento y cambio se persisten atómicamente.", "ENTAILED"),
        ("Negación explícita", "El límite es 1000 operaciones diarias.", "El límite no es 1000 operaciones diarias.", "CONTRADICTED"),
        ("Negación implícita", "Solo administradores eliminan usuarios.", "Un operador puede eliminar usuarios.", "CONTRADICTED"),
        ("Contradicción compleja", "Un fallo deja pendiente y habilita reintento.", "Un fallo confirma y elimina el reintento.", "CONTRADICTED"),
        ("Ausencia", "El sistema utiliza PostgreSQL.", "Soporta exactamente 25.000 concurrentes.", "NOT_ENOUGH_INFORMATION")], [1.15, 2.0, 2.15, 1.2])
    heading(doc, "6. Demostración del límite léxico", 1)
    body(doc, "La evidencia «El límite de transferencias es 1000 operaciones por día» y la respuesta «El límite de transferencias no es 1000 operaciones por día» comparten casi todos sus tokens. SimpleEvaluator puede asignar faithfulness elevada, aunque la partícula «no» invierte la verdad. El test conserva este falso positivo deliberado como prueba de que overlap no equivale a entailment.")
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
    shade_dummy = None
    r=p.add_run("Resultado exigido: el juez semántico debe emitir CONTRADICTED, independientemente de que modelConfidence sea 0.10 o 0.99.")
    set_font(r, 11, True, RED)
    add_figure(doc, diagrams[2], "Figura 3. Flujo de decisión semántica y calibración posterior.")

    heading(doc, "7. Métricas y matriz de confusión", 1)
    body(doc, "La matriz utiliza la etiqueta humana como fila y la predicción del juez como columna. Las métricas se calculan por clase; de este modo, una accuracy aceptable no oculta un bajo recall de contradicciones, que constituye un riesgo crítico para respuestas RAG.")
    add_table(doc, ["Métrica", "Definición", "Uso"], [
        ("Accuracy", "Predicciones correctas / total.", "Visión global; insuficiente ante clases desbalanceadas."),
        ("Precision", "TP / todas las predicciones de la etiqueta.", "Controla falsos positivos del juez."),
        ("Recall", "TP / todos los casos humanos de la etiqueta.", "Controla contradicciones o respaldos omitidos."),
        ("F1", "Media armónica de precision y recall.", "Resume equilibrio por etiqueta."),
        ("Matriz", "Conteos humano × predicho.", "Permite diagnosticar el tipo exacto de error.")], [1.1, 2.35, 3.05])
    heading(doc, "8. Test unitario del benchmark", 1)
    body(doc, "SemanticJudgeBenchmarkTest utiliza seis ejemplos controlados: tres clasificaciones correctas y tres errores deliberados. Verifica accuracy = 0,50; precision de CONTRADICTED = 1/3; recall = 0,50; F1 = 0,40; y las celdas relevantes de la matriz.")
    bullet(doc, "El constructor rechaza un juez nulo.")
    bullet(doc, "evaluate rechaza corpus nulo o vacío.")
    bullet(doc, "Una etiqueta predicha nula produce error inmediato.")
    bullet(doc, "Las vistas externa e interna de la matriz son inmutables.")
    bullet(doc, "Una clase sin casos ni predicciones devuelve precision, recall y F1 iguales a cero.")

    heading(doc, "9. Confianza, certeza y calibración", 1)
    body(doc, "modelConfidence describe cuánto confía el modelo en su propia salida; no demuestra que la evidencia la respalde. Puede ser alta para una afirmación falsa o baja para una verdadera. Por ello, el contrato SemanticJudge no recibe este valor y el benchmark se apoya en etiquetas humanas.")
    add_table(doc, ["Señal", "¿Demuestra verdad?", "Tratamiento"], [
        ("modelConfidence", "No", "Señal secundaria; nunca veredicto autónomo."),
        ("Overlap léxico", "No", "Útil para regresión determinística y detección superficial."),
        ("Etiqueta del juez", "No por sí sola", "Debe medirse contra gold set humano."),
        ("Etiqueta humana", "Referencia operativa", "Debe tener guía, revisión y control de acuerdo."),
        ("Calibración", "Aproxima probabilidad empírica", "Requiere Brier score/ECE y corpus representativo.")], [1.45, 1.45, 3.6])
    heading(doc, "10. Estrategia de ejecución", 1)
    body(doc, "La suite determinística debe ejecutarse en CI sin Ollama. La integración real con Ollama debe mantenerse bajo Maven Failsafe y el perfil ollama-it, evitando tests silenciosamente omitidos.")
    p=doc.add_paragraph(); p.style="Intense Quote"
    set_font(p.add_run("CI determinístico: mvn clean verify\nIntegración Ollama: OLLAMA_BASE_URL=http://localhost:11434 OLLAMA_CHAT_MODEL=qwen3:8b mvn -Pollama-it clean verify"), 10, False, BLUE, "Consolas")
    heading(doc, "11. Criterios de aceptación recomendados", 1)
    bullet(doc, "Corpus versionado, revisado por al menos dos personas para casos ambiguos.")
    bullet(doc, "Recall de CONTRADICTED priorizado por su impacto de riesgo.")
    bullet(doc, "Reporte separado por dominio, idioma y complejidad lingüística.")
    bullet(doc, "Cero dependencia de modelConfidence para asignar la etiqueta semántica.")
    bullet(doc, "Sin regresiones en mvn clean verify y sin pruebas omitidas.")
    bullet(doc, "Recalibración al cambiar modelo, prompt del juez o estrategia RAG.")
    heading(doc, "12. Limitaciones y próximos pasos", 1)
    body(doc, "El juez empleado por las pruebas determinísticas es un doble controlado y no demuestra la calidad de un modelo real. El siguiente paso es implementar un adaptador Ollama para SemanticJudge, exigir salida estructurada, registrar evidencia y decisión, y ejecutar el mismo corpus mediante Failsafe. Los resultados deben compararse por versión del modelo y revisarse ante degradaciones.")
    body(doc, "También se recomienda medir acuerdo entre anotadores —por ejemplo, Cohen’s kappa—, ampliar el corpus con contradicciones temporales, cuantificadores, doble negación y referencias pronominales, y definir umbrales de promoción para CI antes de usar el juez como guardrail productivo.")

    heading(doc, "Anexo A. Trazabilidad", 1)
    add_table(doc, ["Elemento", "Ubicación"], [
        ("Especificación", "SPEC-HALLUCINATION-TESTS.md"),
        ("Prueba semántica", "agent-infrastructure/.../SemanticTruthTest.java"),
        ("Prueba del evaluador", "agent-infrastructure/.../SemanticJudgeBenchmarkTest.java"),
        ("Benchmark", "agent-infrastructure/.../SemanticJudgeBenchmark.java"),
        ("Repositorio", "https://github.com/Prueba008/agente-ollama-spring-ai")], [2.0, 4.5])

    path = OUT / "Validacion_Verdad_Semantica_Agente_Ollama.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    diagrams = [diagram_architecture(), diagram_sequence(), diagram_flow()]
    print(build_doc(diagrams))
